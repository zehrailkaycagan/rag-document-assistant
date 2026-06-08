"""
Doküman yükleme modülü - PDF, DOCX, TXT dosyalarını okur
"""
from typing import List, Optional
from pathlib import Path
import logging

from pypdf import PdfReader
from docx import Document

from src.utils import getFileExtension, cleanText

logger = logging.getLogger(__name__)

# pdfplumber'ı opsiyonel olarak yükle (tablo okuma için)
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber yüklü değil, tablo okuma özelliği devre dışı. 'pip install pdfplumber' ile yükleyebilirsiniz.")

class DocumentLoader:
    """Farklı formatlardaki dokümanları yükler"""
    
    @staticmethod
    def _formatTableAsText(table: List[List]) -> str:
        """Tabloyu okunabilir metin formatına çevirir"""
        if not table:
            return ""
        
        lines = []
        for row in table:
            if row:
                # Satırı pipe ile ayır ve boş hücreleri temizle
                formattedRow = " | ".join([str(cell).strip() if cell else "" for cell in row])
                if formattedRow.strip():  # Boş satırları atla
                    lines.append(formattedRow)
        
        if lines:
            return "\n[TABLO]\n" + "\n".join(lines) + "\n[TABLO SON]\n"
        return ""
    
    @staticmethod
    def _extractTablesWithPdfplumber(filePath: str) -> str:
        """pdfplumber kullanarak PDF'den tabloları çıkarır"""
        if not PDFPLUMBER_AVAILABLE:
            return ""
        
        tableTexts = []
        try:
            with pdfplumber.open(filePath) as pdf:
                for pageNum, page in enumerate(pdf.pages, 1):
                    try:
                        tables = page.extract_tables()
                        if tables:
                            for tableNum, table in enumerate(tables, 1):
                                if table:
                                    formattedTable = DocumentLoader._formatTableAsText(table)
                                    if formattedTable:
                                        tableTexts.append(f"[Sayfa {pageNum}, Tablo {tableNum}]\n{formattedTable}")
                                        logger.debug(f"Sayfa {pageNum}'den tablo {tableNum} çıkarıldı")
                    except Exception as pageError:
                        logger.warning(f"Sayfa {pageNum} tabloları okunamadı: {pageError}")
                        continue
        except Exception as e:
            logger.warning(f"pdfplumber ile tablo okuma hatası: {e}")
            return ""
        
        if tableTexts:
            logger.info(f"Toplam {len(tableTexts)} tablo çıkarıldı")
            return "\n\n".join(tableTexts)
        return ""
    
    @staticmethod
    def loadPDF(filePath: str) -> str:
        """PDF dosyasını okur ve metne çevirir"""
        try:
            reader = PdfReader(filePath, strict=False)
            textParts = []
            
            # PDF sayfa sayısını kontrol et
            if len(reader.pages) == 0:
                logger.warning(f"PDF dosyası boş: {filePath}")
                return ""
            
            # Şifre kontrolü
            if reader.is_encrypted:
                logger.warning(f"PDF şifreli, şifresiz okuma deneniyor: {filePath}")
                try:
                    reader.decrypt("")
                except Exception:
                    raise ValueError("PDF dosyası şifreli ve şifre bilinmiyor")
            
            # Sayfaları oku
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        textParts.append(text)
                except Exception as pageError:
                    logger.warning(f"PDF sayfa {i+1} okunamadı: {pageError}")
                    continue
            
            # Tabloları çıkar (pdfplumber ile)
            tableText = ""
            if PDFPLUMBER_AVAILABLE:
                try:
                    tableText = DocumentLoader._extractTablesWithPdfplumber(filePath)
                    if tableText:
                        logger.info(f"PDF'den tablolar çıkarıldı: {filePath}")
                except Exception as tableError:
                    logger.warning(f"Tablolar çıkarılamadı: {tableError}")
            
            # Metin ve tabloları birleştir
            if not textParts and not tableText:
                logger.warning(f"PDF'den metin veya tablo çıkarılamadı: {filePath}")
                return ""
            
            # Metinleri birleştir
            fullText = "\n".join(textParts) if textParts else ""
            
            # Tabloları ekle
            if tableText:
                if fullText:
                    fullText = fullText + "\n\n" + tableText
                else:
                    fullText = tableText
            
            cleanedText = cleanText(fullText)
            
            if not cleanedText or len(cleanedText.strip()) == 0:
                logger.warning(f"PDF temizlendikten sonra boş kaldı: {filePath}")
                return ""
            
            return cleanedText
            
        except ValueError as ve:
            # Zaten işlenmiş hata mesajları
            raise ve
        except Exception as e:
            errorMsg = str(e)
            logger.error(f"PDF okuma hatası ({filePath}): {errorMsg}")
            
            # Özel hata mesajları
            if "encrypted" in errorMsg.lower() or "password" in errorMsg.lower():
                raise ValueError("PDF dosyası şifreli. Lütfen şifresiz bir PDF kullanın.")
            elif "corrupt" in errorMsg.lower() or "damaged" in errorMsg.lower():
                raise ValueError("PDF dosyası bozuk veya hasarlı.")
            else:
                raise ValueError(f"PDF dosyası okunamadı: {errorMsg}")
    
    @staticmethod
    def loadDOCX(filePath: str) -> str:
        """DOCX dosyasını okur ve metne çevirir (tablolar dahil)"""
        try:
            doc = Document(filePath)
            textParts = []
            
            # Paragrafları oku
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    textParts.append(paragraph.text)
            
            # Tabloları oku
            for tableNum, table in enumerate(doc.tables, 1):
                try:
                    tableText = DocumentLoader._formatTableAsText(
                        [[cell.text.strip() if cell.text else "" for cell in row.cells] for row in table.rows]
                    )
                    if tableText:
                        textParts.append(f"[Tablo {tableNum}]\n{tableText}")
                        logger.debug(f"DOCX'den tablo {tableNum} çıkarıldı")
                except Exception as tableError:
                    logger.warning(f"DOCX tablo {tableNum} okunamadı: {tableError}")
                    continue
            
            fullText = "\n".join(textParts)
            return cleanText(fullText)
        except Exception as e:
            logger.error(f"DOCX okuma hatası: {e}")
            raise ValueError(f"DOCX dosyası okunamadı: {str(e)}")
    
    @staticmethod
    def loadTXT(filePath: str) -> str:
        """TXT dosyasını okur"""
        try:
            with open(filePath, "r", encoding="utf-8") as f:
                text = f.read()
            return cleanText(text)
        except Exception as e:
            logger.error(f"TXT okuma hatası: {e}")
            raise ValueError(f"TXT dosyası okunamadı: {str(e)}")
    
    @classmethod
    def loadDocument(cls, filePath: str) -> str:
        """Dosya tipine göre uygun yükleme metodunu çağırır"""
        fileExtension = getFileExtension(filePath)
        
        if not Path(filePath).exists():
            raise FileNotFoundError(f"Dosya bulunamadı: {filePath}")
        
        if fileExtension == ".pdf":
            return cls.loadPDF(filePath)
        elif fileExtension == ".docx":
            return cls.loadDOCX(filePath)
        elif fileExtension == ".txt":
            return cls.loadTXT(filePath)
        else:
            raise ValueError(f"Desteklenmeyen dosya formatı: {fileExtension}")
    
    @classmethod
    def loadMultipleDocuments(cls, filePaths: List[str]) -> List[dict]:
        """Birden fazla dokümanı yükler"""
        documents = []
        failedFiles = []
        
        for filePath in filePaths:
            try:
                text = cls.loadDocument(filePath)
                
                # Boş metin kontrolü
                if not text or not text.strip():
                    logger.warning(f"Dosya boş veya metin çıkarılamadı: {filePath}")
                    failedFiles.append((Path(filePath).name, "Dosya boş veya metin çıkarılamadı"))
                    continue
                
                documents.append({
                    "filePath": filePath,
                    "fileName": Path(filePath).name,
                    "text": text,
                    "fileType": getFileExtension(filePath)
                })
                logger.info(f"Dosya başarıyla yüklendi: {Path(filePath).name}")
                
            except FileNotFoundError as fnfe:
                errorMsg = f"Dosya bulunamadı: {filePath}"
                logger.error(errorMsg)
                failedFiles.append((Path(filePath).name, str(fnfe)))
            except ValueError as ve:
                errorMsg = str(ve)
                logger.error(f"Dosya yükleme hatası ({filePath}): {errorMsg}")
                failedFiles.append((Path(filePath).name, errorMsg))
            except Exception as e:
                errorMsg = f"Beklenmeyen hata: {str(e)}"
                logger.error(f"Dosya yüklenemedi {filePath}: {errorMsg}")
                failedFiles.append((Path(filePath).name, errorMsg))
        
        if failedFiles:
            failedNames = ", ".join([name for name, _ in failedFiles])
            logger.warning(f"Başarısız dosyalar: {failedNames}")
        
        if not documents and failedFiles:
            errorDetails = "; ".join([f"{name}: {reason}" for name, reason in failedFiles])
            raise ValueError(f"Hiçbir dosya yüklenemedi. Hatalar: {errorDetails}")
        
        return documents
